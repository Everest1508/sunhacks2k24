import PyPDF2, pdfplumber
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from django.conf import settings
import cv2
import pytesseract as tess
import numpy as np
import os
tess.pytesseract.tesseract_cmd=r'C:\Program Files\Tesseract-OCR\tesseract.exe'
from docx import Document

def calculate_score(cv_file,jd_file):
    if cv_file.name.split(".")[-1]=="pdf":
        Script=PyPDF2.PdfReader(cv_file)
        pages=len(Script.pages)
        Script = []
        with pdfplumber.open(cv_file) as pdf:
            for i in range (0,pages):
                page=pdf.pages[i]
                text=page.extract_text()
                print (text)
                Script.append(text)
        Script=''.join(Script)
    elif cv_file.name.split(".")[-1]=="docs" or cv_file.name.split(".")[-1]=="docx" or cv_file.name.split(".")[-1]=="doc" :
        doc = Document(cv_file.path)
        content = [p.text for p in doc.paragraphs]
        Script = " ".join(content)
    else:
        img = cv2.imread(cv_file.path)
        print(cv_file.url)
        gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
        gray, img_bin = cv2.threshold(gray,128,255,cv2.THRESH_BINARY | cv2.THRESH_OTSU)
        gray = cv2.bitwise_not(img_bin)
        kernel = np.ones((2, 1), np.uint8)
        img = cv2.erode(gray, kernel, iterations=1)
        img = cv2.dilate(img, kernel, iterations=1)
        out_below = tess.image_to_string(img)
        Script = out_below
        
    CV_Clear=Script.replace("\n","")

    Script_Req=jd_file
    Req_Clear=Script_Req.replace("\n","")
    Req_Clear
    Match_Test=[CV_Clear,Req_Clear]
    cv=CountVectorizer()
    count_matrix=cv.fit_transform(Match_Test)
    print('Similarity is :',cosine_similarity(count_matrix))
    MatchPercentage=cosine_similarity(count_matrix)[0][1]*100
    MatchPercentage=round(MatchPercentage,2)
    print('Match Percentage is :'+ str(MatchPercentage)+'% to Requirement')
    print(type(MatchPercentage))
    return MatchPercentage

def extract_text(cv_file):
    if cv_file.name.split(".")[-1]=="pdf":
        Script=PyPDF2.PdfReader(cv_file)
        pages=len(Script.pages)
        Script = []
        with pdfplumber.open(cv_file) as pdf:
            for i in range (0,pages):
                page=pdf.pages[i]
                text=page.extract_text()
                print (text)
                Script.append(text)
        Script=''.join(Script)
    elif cv_file.name.split(".")[-1]=="docs" or cv_file.name.split(".")[-1]=="docx" or cv_file.name.split(".")[-1]=="doc" :
        doc = Document(cv_file)
        content = [p.text for p in doc.paragraphs]
        Script = " ".join(content)
    else:
        nparr = np.frombuffer(cv_file.read(), np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        # img = cv2.imread(cv_file)
        # print(cv_file.url)
        gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
        gray, img_bin = cv2.threshold(gray,128,255,cv2.THRESH_BINARY | cv2.THRESH_OTSU)
        gray = cv2.bitwise_not(img_bin)
        kernel = np.ones((2, 1), np.uint8)
        img = cv2.erode(gray, kernel, iterations=1)
        img = cv2.dilate(img, kernel, iterations=1)
        out_below = tess.image_to_string(img)
        Script = out_below
        
    return Script.replace("\n","")